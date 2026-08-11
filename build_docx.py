"""Build an editable two-column paper.docx from `paper.md`.

Layout mirrors the PDF: title + author + abstract span the full page width;
the body (from section 1 onward) flows in two columns; every wide table and the
figure sits in its own full-width (single-column) section so it stays readable.

Run: .venv/Scripts/python.exe build_docx.py
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
MD = ROOT / "paper.md"
FIG = ROOT / "figures" / "figure1_survival.png"
OUT = ROOT / "paper.docx"

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+?`|\*[^*]+?\*)")


def add_runs(p, text, base=10.0):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            r = p.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(base * 0.9)
        elif part.startswith("*") and part.endswith("*") and len(part) >= 2:
            r = p.add_run(part[1:-1]); r.italic = True
        else:
            p.add_run(part)


# ---------- markdown -> block list ----------
def is_tbl(s):
    s = s.strip()
    return s.startswith("|") and s.endswith("|")


def parse_blocks(lines):
    items, i, n = [], 0, len(lines)
    while i < n:
        ln = lines[i]
        s = ln.strip()
        if s == "":
            i += 1; continue
        if ln.startswith("> "):                       # blockquote group (full-width)
            inner = []
            while i < n and (lines[i].startswith(">")):
                inner.append(lines[i][2:] if lines[i].startswith("> ") else lines[i][1:])
                i += 1
            items.append(("group", parse_blocks(inner)))
            continue
        if is_tbl(ln):                                # bare table (full-width)
            rows = []
            while i < n and is_tbl(lines[i]):
                rows.append(lines[i]); i += 1
            items.append(("table", parse_table(rows)))
            continue
        if ln.startswith("# "):
            items.append(("h1", ln[2:].strip())); i += 1; continue
        if ln.startswith("### "):
            items.append(("h3", ln[4:].strip())); i += 1; continue
        if ln.startswith("## "):
            items.append(("h2", ln[3:].strip())); i += 1; continue
        if ln.strip() == "---":
            items.append(("hr",)); i += 1; continue
        m = re.match(r"^(\d+)\.\s+(.*)", ln)
        if m:
            block = []
            while i < n and re.match(r"^\d+\.\s+", lines[i]):
                block.append(re.match(r"^\d+\.\s+(.*)", lines[i]).group(1)); i += 1
            items.append(("numbers", block)); continue
        if ln.startswith("- "):
            block = []
            while i < n and lines[i].startswith("- "):
                block.append(lines[i][2:]); i += 1
            items.append(("bullets", block)); continue
        items.append(("p", ln)); i += 1
    return items


def parse_table(rows):
    out = []
    for idx, r in enumerate(rows):
        if idx == 1:                                  # separator row
            continue
        cells = [c.strip() for c in r.strip().strip("|").split("|")]
        out.append(cells)
    return out


# ---------- section / column helpers ----------
def setup_page(sec):
    sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
    sec.left_margin = Cm(1.4); sec.right_margin = Cm(1.4)
    sec.top_margin = Cm(1.5); sec.bottom_margin = Cm(1.6)


def set_cols(sec, num):
    sectPr = sec._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols"); sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), "340")
    cols.set(qn("w:equalWidth"), "1")


# ---------- renderers ----------
def render_para(doc, text, size=10.0, bold=False, italic=False, align=None,
                space_before=0.0, space_after=4.0, keep_next=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before); pf.space_after = Pt(space_after)
    pf.line_spacing = 1.12
    if align is not None:
        p.alignment = align
    if keep_next:
        pf.keep_with_next = True
    if bold or italic:
        r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    else:
        add_runs(p, text, base=size)
    for r in p.runs:
        if r.font.size is None:
            r.font.size = Pt(size)
    return p


def render_heading(doc, text, level):
    if level == 1:
        render_para(doc, text, size=17, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before=2, space_after=6, keep_next=True)
    elif level == 2:
        render_para(doc, text, size=12.5, bold=True, space_before=9, space_after=4, keep_next=True)
    else:
        render_para(doc, text, size=10.5, bold=True, space_before=7, space_after=3, keep_next=True)


def render_list(doc, texts, numbered):
    style = "List Number" if numbered else "List Bullet"
    for t in texts:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.1
        add_runs(p, t, base=10)
        for r in p.runs:
            if r.font.size is None:
                r.font.size = Pt(10)


def render_table(doc, rows):
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncol)
    t.style = "Table Grid"
    t.alignment = 1  # center
    for ri, row in enumerate(rows):
        for ci in range(ncol):
            cell = t.cell(ri, ci)
            cell.paragraphs[0].paragraph_format.space_after = Pt(1)
            cell.paragraphs[0].paragraph_format.line_spacing = 1.0
            txt = row[ci] if ci < len(row) else ""
            add_runs(cell.paragraphs[0], txt, base=8)
            for r in cell.paragraphs[0].runs:
                if r.font.size is None:
                    r.font.size = Pt(8)
                if ri == 0:
                    r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def render_image(doc, path):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    if path.exists():
        p.add_run().add_picture(str(path), width=Cm(16.5))
    else:
        p.add_run("[figure missing]")


def render_group(doc, sub_items):
    """A blockquote group: caption(s), table/figure, notes — all full width."""
    for it in sub_items:
        k = it[0]
        if k == "table":
            render_table(doc, it[1])
        elif k == "p":
            txt = it[1]
            if "figure1_survival.png" in txt:
                render_image(doc, FIG)
            elif txt.strip().startswith("**"):
                render_para(doc, txt, size=9.5, space_before=3, space_after=2, keep_next=True)
            else:
                render_para(doc, txt, size=8.8, space_after=2)
        elif k == "group":
            render_group(doc, it[1])


# ---------- main build ----------
def main():
    text = MD.read_text(encoding="utf-8")
    items = parse_blocks(text.split("\n"))

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"; normal.font.size = Pt(10)

    setup_page(doc.sections[0]); set_cols(doc.sections[0], 1)
    state = {"cols": 1}

    def switch(cols):
        if state["cols"] != cols:
            sec = doc.add_section(WD_SECTION.CONTINUOUS)
            setup_page(sec); set_cols(sec, cols)
            state["cols"] = cols

    body_started = False
    FULLW = {"table", "group"}

    for it in items:
        kind = it[0]
        is_body_h2 = (kind == "h2" and it[1].strip().startswith("1."))
        if is_body_h2 and not body_started:
            body_started = True
            switch(2)

        if not body_started:                          # title block, full width
            if kind == "h1":
                render_heading(doc, it[1], 1)
            elif kind == "h2":
                render_heading(doc, it[1], 2)
            elif kind == "p":
                t = it[1]
                if "Aqrawi" in t and t.strip().startswith("**"):
                    render_para(doc, t.strip("*"), size=12, bold=True,
                                align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
                else:
                    render_para(doc, t, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            continue

        # body
        if kind in FULLW:
            switch(1)
            if kind == "table":
                render_table(doc, it[1])
            else:
                render_group(doc, it[1])
        elif kind == "hr":
            continue
        else:
            switch(2)
            if kind == "h2":
                render_heading(doc, it[1], 2)
            elif kind == "h3":
                render_heading(doc, it[1], 3)
            elif kind == "numbers":
                render_list(doc, it[1], numbered=True)
            elif kind == "bullets":
                render_list(doc, it[1], numbered=False)
            elif kind == "p":
                render_para(doc, it[1], size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    doc.save(OUT)

    # ---- verify ----
    chk = Document(str(OUT))
    ncols = []
    for s in chk.sections:
        c = s._sectPr.find(qn("w:cols"))
        ncols.append(c.get(qn("w:num")) if c is not None else "1")
    print(f"saved {OUT.name}: {len(chk.paragraphs)} paragraphs, "
          f"{len(chk.tables)} tables, {len(chk.sections)} sections")
    print("section column counts in order:", ncols)
    print("em dashes in docx text:",
          sum(p.text.count('—') for p in chk.paragraphs))


if __name__ == "__main__":
    main()
